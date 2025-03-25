import os
import pandas as pd
import json
import logging
from PyPDF2 import PdfReader
from docx import Document
import openpyxl
from models.db import db
from models.file import File
from services.ai_service import summarize_text
from typing import Dict, Any, Optional
from datetime import datetime
from pathlib import Path
from utils.validation import validate_file_path, validate_file_extension
from utils.retry import retry, RetryContext
from utils.decorators import circuit_breaker, CircuitBreaker
from utils.exceptions import (
    FileProcessingError, ValidationError, ResourceCleanupError,
    ExternalServiceError
)
import magic
import mimetypes
from werkzeug.utils import secure_filename

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try to use python-magic-bin, fallback to mimetypes
try:
    def get_mime_type(file_path):
        return magic.from_file(file_path, mime=True)
except (ImportError, OSError):
    def get_mime_type(file_path):
        return mimetypes.guess_type(file_path)[0] or 'application/octet-stream'

class FileService:
    """Service for handling file operations"""
    
    def __init__(self, upload_folder: str = 'uploads'):
        self.upload_folder = upload_folder
        self.supported_formats = {
            'csv': self._load_csv,
            'xlsx': self._load_excel,
            'json': self._load_json,
            'txt': self._load_text
        }
        self.circuit_breaker = CircuitBreaker(
            failure_threshold=3,
            reset_timeout=30.0
        )
    
    @retry(max_attempts=3, delay=1.0, backoff=2.0)
    def process_file(self, file_id: str, file_path: str, file_extension: str) -> None:
        """Process file based on its type"""
        file = None
        try:
            # Validate file path
            validate_file_path(file_path)
            validate_file_extension(file_path, list(self.supported_formats.keys()))
            
            # Get file from database
            file = self._get_file_from_db(file_id)
            if not file:
                raise FileProcessingError(f"File not found: {file_id}")
            
            # Update processing status
            file.processing_status = 'processing'
            self._commit_db_changes()
            
            # Process based on file type
            if file_extension in ['pdf', 'doc', 'docx', 'txt', 'md']:
                self._process_text_document(file, file_path)
            
            elif file_extension in ['csv', 'xls', 'xlsx']:
                self._process_tabular_data(file, file_path)
            
            elif file_extension in ['json', 'xml']:
                self._process_structured_data(file, file_path)
            
            elif file_extension in ['jpg', 'jpeg', 'png', 'gif']:
                self._process_image(file, file_path)
            
            # Update processing status
            file.is_processed = True
            file.processing_status = 'completed'
            self._commit_db_changes()
        
        except ValidationError as e:
            logger.error(f"Validation error processing file: {str(e)}")
            self._update_file_status(file, 'failed', str(e))
            raise
        except ExternalServiceError as e:
            logger.error(f"External service error processing file: {str(e)}")
            self._update_file_status(file, 'failed', str(e))
            raise
        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            self._update_file_status(file, 'failed', str(e))
            raise FileProcessingError(f"Failed to process file: {str(e)}")
        finally:
            self._cleanup_resources()
    
    @retry(max_attempts=3, delay=1.0, backoff=2.0)
    @circuit_breaker
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from file"""
        try:
            # Validate file path
            validate_file_path(file_path)
            
            file_extension = Path(file_path).suffix.lower().replace('.', '')
            
            if file_extension in ['txt', 'md']:
                return self._extract_text_metadata(file_path)
            elif file_extension == 'json':
                return self._extract_json_metadata(file_path)
            elif file_extension in ['jpg', 'jpeg', 'png', 'gif']:
                return self._extract_image_metadata(file_path)
            else:
                raise ValidationError(f"Unsupported file type for metadata extraction: {file_extension}")
        
        except ValidationError as e:
            logger.error(f"Validation error extracting metadata: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            raise FileProcessingError(f"Failed to extract metadata: {str(e)}")
    
    def _extract_text_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from text file"""
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                lines = file.readlines()
                
                metadata['line_count'] = len(lines)
                metadata['word_count'] = sum(len(line.split()) for line in lines)
                
                # Generate summary
                text = ''.join(lines[:20])
                metadata['summary'] = summarize_text(text, 100)
        
        except Exception as e:
            logger.error(f"Error extracting text metadata: {str(e)}")
            raise FileProcessingError(f"Failed to extract text metadata: {str(e)}")
        
        return metadata
    
    def _extract_json_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from JSON file"""
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                data = json.load(file)
                
                # Generate summary based on keys
                if isinstance(data, dict):
                    keys = list(data.keys())
                    metadata['summary'] = f"JSON file with keys: {', '.join(keys[:5])}"
                    if len(keys) > 5:
                        metadata['summary'] += f" and {len(keys) - 5} more keys"
                
                elif isinstance(data, list):
                    metadata['summary'] = f"JSON array with {len(data)} items"
                    if data and isinstance(data[0], dict):
                        keys = list(data[0].keys())
                        metadata['summary'] += f", first item has keys: {', '.join(keys[:5])}"
                        if len(keys) > 5:
                            metadata['summary'] += f" and {len(keys) - 5} more keys"
        
        except Exception as e:
            logger.error(f"Error extracting JSON metadata: {str(e)}")
            raise FileProcessingError(f"Failed to extract JSON metadata: {str(e)}")
        
        return metadata
    
    def _extract_image_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from image file"""
        metadata = {
            'has_images': True
        }
        
        try:
            from PIL import Image
            
            with Image.open(file_path) as img:
                metadata['width'] = img.width
                metadata['height'] = img.height
                metadata['format'] = img.format
                metadata['mode'] = img.mode
                metadata['summary'] = f"{img.format} image, {img.width}x{img.height} pixels, {img.mode} mode"
        
        except Exception as e:
            logger.error(f"Error extracting image metadata: {str(e)}")
            raise FileProcessingError(f"Failed to extract image metadata: {str(e)}")
        
        return metadata
    
    def _update_file_status(self, file: Any, status: str, error: Optional[str] = None) -> None:
        """Update file processing status"""
        try:
            if file:
                file.processing_status = status
                if error:
                    file.processing_error = error
                self._commit_db_changes()
        except Exception as e:
            logger.error(f"Error updating file status: {str(e)}")
            raise FileProcessingError(f"Failed to update file status: {str(e)}")
    
    def _cleanup_resources(self) -> None:
        """Clean up resources"""
        try:
            # Add any cleanup logic here
            pass
        except Exception as e:
            logger.error(f"Error during cleanup: {str(e)}")
            raise ResourceCleanupError(f"Failed to clean up resources: {str(e)}")
    
    def __enter__(self):
        """Context manager entry"""
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit with cleanup"""
        self._cleanup_resources()

def get_file_metadata(file_path, file_extension):
    """Extract metadata from file based on its type"""
    metadata = {
        'page_count': None,
        'word_count': None,
        'line_count': None,
        'column_count': None,
        'row_count': None,
        'sheet_count': None,
        'has_images': False,
        'has_tables': False,
        'has_charts': False,
        'summary': None
    }
    
    try:
        # Process based on file type
        if file_extension in ['pdf']:
            metadata.update(extract_pdf_metadata(file_path))
        
        elif file_extension in ['doc', 'docx']:
            metadata.update(extract_docx_metadata(file_path))
        
        elif file_extension in ['xls', 'xlsx']:
            metadata.update(extract_excel_metadata(file_path))
        
        elif file_extension in ['csv']:
            metadata.update(extract_csv_metadata(file_path))
        
        elif file_extension in ['txt', 'md']:
            metadata.update(extract_text_metadata(file_path))
        
        elif file_extension in ['json']:
            metadata.update(extract_json_metadata(file_path))
        
        elif file_extension in ['jpg', 'jpeg', 'png', 'gif']:
            metadata.update(extract_image_metadata(file_path))
    
    except Exception as e:
        logger.error(f"Error extracting metadata: {str(e)}")
    
    return metadata


def extract_pdf_metadata(file_path):
    """Extract metadata from PDF file"""
    metadata = {}
    
    try:
        with open(file_path, 'rb') as file:
            pdf = PdfReader(file)
            
            # Get page count
            metadata['page_count'] = len(pdf.pages)
            
            # Extract text from first page for summary
            if len(pdf.pages) > 0:
                first_page_text = pdf.pages[0].extract_text()
                metadata['summary'] = summarize_text(first_page_text, 100)
            
            # Check for images (basic check)
            for page in pdf.pages[:min(5, len(pdf.pages))]:
                if '/XObject' in page:
                    metadata['has_images'] = True
                    break
    
    except Exception as e:
        logger.error(f"Error extracting PDF metadata: {str(e)}")
    
    return metadata


def extract_docx_metadata(file_path):
    """Extract metadata from DOCX file"""
    metadata = {}
    
    try:
        doc = Document(file_path)
        
        # Count paragraphs, tables, and words
        metadata['page_count'] = len(doc.paragraphs) // 40 + 1  # Rough estimate
        metadata['word_count'] = sum(len(p.text.split()) for p in doc.paragraphs)
        metadata['line_count'] = len(doc.paragraphs)
        metadata['has_tables'] = len(doc.tables) > 0
        
        # Extract text for summary
        text = '\n'.join(p.text for p in doc.paragraphs[:10])
        metadata['summary'] = summarize_text(text, 100)
    
    except Exception as e:
        logger.error(f"Error extracting DOCX metadata: {str(e)}")
    
    return metadata


def extract_excel_metadata(file_path):
    """Extract metadata from Excel file"""
    metadata = {}
    
    try:
        workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        
        # Count sheets
        metadata['sheet_count'] = len(workbook.sheetnames)
        
        # Get data from first sheet
        if workbook.sheetnames:
            sheet = workbook[workbook.sheetnames[0]]
            metadata['row_count'] = sheet.max_row
            metadata['column_count'] = sheet.max_column
            
            # Check for charts
            metadata['has_charts'] = hasattr(sheet, '_charts') and len(sheet._charts) > 0
    
    except Exception as e:
        logger.error(f"Error extracting Excel metadata: {str(e)}")
    
    return metadata


def extract_csv_metadata(file_path):
    """Extract metadata from CSV file"""
    metadata = {}
    
    try:
        df = pd.read_csv(file_path)
        
        metadata['row_count'] = len(df)
        metadata['column_count'] = len(df.columns)
        metadata['has_tables'] = True
        
        # Generate summary from column names
        metadata['summary'] = f"CSV file with {len(df)} rows and {len(df.columns)} columns: {', '.join(df.columns[:5])}"
        if len(df.columns) > 5:
            metadata['summary'] += f" and {len(df.columns) - 5} more columns"
    
    except Exception as e:
        logger.error(f"Error extracting CSV metadata: {str(e)}")
    
    return metadata


def extract_text_metadata(file_path):
    """Extract metadata from text file"""
    metadata = {}
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            lines = file.readlines()
            
            metadata['line_count'] = len(lines)
            metadata['word_count'] = sum(len(line.split()) for line in lines)
            
            # Generate summary
            text = ''.join(lines[:20])
            metadata['summary'] = summarize_text(text, 100)
    
    except Exception as e:
        logger.error(f"Error extracting text metadata: {str(e)}")
    
    return metadata


def extract_json_metadata(file_path):
    """Extract metadata from JSON file"""
    metadata = {}
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            data = json.load(file)
            
            # Generate summary based on keys
            if isinstance(data, dict):
                keys = list(data.keys())
                metadata['summary'] = f"JSON file with keys: {', '.join(keys[:5])}"
                if len(keys) > 5:
                    metadata['summary'] += f" and {len(keys) - 5} more keys"
            
            elif isinstance(data, list):
                metadata['summary'] = f"JSON array with {len(data)} items"
                if data and isinstance(data[0], dict):
                    keys = list(data[0].keys())
                    metadata['summary'] += f", first item has keys: {', '.join(keys[:5])}"
                    if len(keys) > 5:
                        metadata['summary'] += f" and {len(keys) - 5} more keys"
    
    except Exception as e:
        logger.error(f"Error extracting JSON metadata: {str(e)}")
    
    return metadata


def extract_image_metadata(file_path):
    """Extract metadata from image file"""
    metadata = {
        'has_images': True
    }
    
    try:
        from PIL import Image
        
        with Image.open(file_path) as img:
            metadata['width'] = img.width
            metadata['height'] = img.height
            metadata['format'] = img.format
            metadata['mode'] = img.mode
            metadata['summary'] = f"{img.format} image, {img.width}x{img.height} pixels, {img.mode} mode"
    
    except Exception as e:
        logger.error(f"Error extracting image metadata: {str(e)}")
    
    return metadata


def process_text_document(file, file_path):
    """Process text document (PDF, DOCX, TXT, etc.)"""
    # In a real application, this would extract text, analyze content, etc.
    logger.info(f"Processing text document: {file.original_filename}")


def process_tabular_data(file, file_path):
    """Process tabular data (CSV, Excel, etc.)"""
    # In a real application, this would analyze data, generate statistics, etc.
    logger.info(f"Processing tabular data: {file.original_filename}")


def process_structured_data(file, file_path):
    """Process structured data (JSON, XML, etc.)"""
    # In a real application, this would parse and analyze the data
    logger.info(f"Processing structured data: {file.original_filename}")


def process_image(file, file_path):
    """Process image file"""
    # In a real application, this would analyze the image, extract features, etc.
    logger.info(f"Processing image: {file.original_filename}") 